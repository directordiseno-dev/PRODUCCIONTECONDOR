from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


OUTPUT = Path(r"C:\Users\auxil\Desktop\PRODUCCION TECONDOR APP\Hoja de vida - Mateo Agudelo.docx")
PHOTO = Path(r"C:\Users\auxil\Desktop\PRODUCCION TECONDOR APP\tmp\cv_original\portrait.png")

FONT = "Arial"
INK = RGBColor(45, 45, 45)
TEXT = RGBColor(99, 99, 99)
SUBTLE = RGBColor(76, 80, 84)
RULE = "5B5B5B"
HEADER_FILL = "F2F2F2"


def set_run_font(run, size, bold=False, color=TEXT, italic=False, spacing=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    if spacing is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_twips))
    tcw.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths):
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "0")
    tblind.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblpr.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        cell.margin_top = 0
        cell.margin_bottom = 0


def add_bottom_border(paragraph, color=RULE, size="5", space="3"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.32)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.0)
    normal.paragraph_format.line_spacing = 1.08

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.0)
    bullet.font.color.rgb = TEXT
    bullet.paragraph_format.left_indent = Inches(0.14)
    bullet.paragraph_format.first_line_indent = Inches(-0.10)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(2.0)
    bullet.paragraph_format.line_spacing = 1.05


def add_header(doc):
    table = doc.add_table(rows=1, cols=2)
    configure_table_geometry(table, [8740, 1870])
    left, right = table.rows[0].cells
    set_cell_shading(left, HEADER_FILL)
    set_cell_shading(right, HEADER_FILL)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("MATEO AGUDELO")
    set_run_font(run, 32, bold=True, color=INK, spacing=75)

    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("I N G E N I E R O   Q U Í M I C O")
    set_run_font(run, 10.5, bold=True, color=SUBTLE)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if PHOTO.exists():
        p.add_run().add_picture(str(PHOTO), width=Inches(1.45))

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.line_spacing = 0.1
    after.add_run(" ")


def start_two_columns(doc):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.32)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    sectpr = section._sectPr
    cols = sectpr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectpr.append(cols)
    for child in list(cols):
        cols.remove(child)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:equalWidth"), "0")
    cols.set(qn("w:sep"), "1")
    cols.set(qn("w:space"), "360")
    left = OxmlElement("w:col")
    left.set(qn("w:w"), "3850")
    left.set(qn("w:space"), "360")
    right = OxmlElement("w:col")
    right.set(qn("w:w"), "6400")
    cols.append(left)
    cols.append(right)
    return section


def add_heading(doc, text, size=14.2, top=9, bottom=6, rule=True):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(top)
    p.paragraph_format.space_after = Pt(bottom)
    if rule:
        add_bottom_border(p)
    run = p.add_run(text.upper())
    set_run_font(run, size, bold=True, color=INK, spacing=55)
    return p


def add_contact(doc, symbol, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.5)
    p.paragraph_format.left_indent = Inches(0.02)
    run = p.add_run(f"{symbol}  ")
    set_run_font(run, 9.2, bold=True, color=INK)
    run = p.add_run(text)
    set_run_font(run, 9.2, color=TEXT)


def add_sidebar_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3.0)
    run = p.add_run(text)
    set_run_font(run, 9.0, color=TEXT)


def add_education(doc, degree, institution, dates):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5.0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(degree.upper())
    set_run_font(run, 9.6, bold=True, color=SUBTLE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(institution)
    set_run_font(run, 9.1, color=TEXT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.0)
    run = p.add_run(dates)
    set_run_font(run, 9.1, color=SUBTLE)


def add_profile(doc):
    add_heading(doc, "Perfil", size=14.2, top=0, bottom=6, rule=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "Ingeniero químico con experiencia en supervisión de procesos productivos, control de calidad, "
        "mejora continua y análisis de datos. Ha liderado personal de planta, estandarizado formulaciones "
        "y controlado variables críticas como pH, viscosidad y sólidos para asegurar el cumplimiento de "
        "especificaciones. Integra producción, calidad y herramientas digitales para identificar desviaciones, "
        "implementar acciones correctivas y mejorar la productividad."
    )
    run = p.add_run(text)
    set_run_font(run, 9.35, color=TEXT)


def add_job(doc, title, company, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5.0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title.upper())
    set_run_font(run, 9.8, bold=True, color=SUBTLE)

    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(company)
    set_run_font(run, 9.2, color=TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2.0)
    run = p.add_run(dates.upper())
    set_run_font(run, 8.9, color=SUBTLE)

    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        run = p.add_run(bullet)
        set_run_font(run, 8.9, color=TEXT)


def build():
    doc = Document()
    configure_document(doc)
    add_header(doc)
    start_two_columns(doc)

    add_contact(doc, "☎", "314 323 7380")
    add_contact(doc, "✉", "matius-098@gmail.com")
    add_contact(doc, "⌖", "Medellín, Antioquia")
    add_contact(doc, "◈", "github.com/Mateoiq98")
    add_contact(doc, "in", "linkedin.com/in/mateoiq98")

    add_heading(doc, "Habilidades", size=14.0, top=9, bottom=6)
    for item in [
        "Supervisión de procesos químicos",
        "Liderazgo de personal operativo",
        "Estandarización de formulaciones",
        "Control de variables críticas",
        "BPM y control de calidad",
        "Mejora continua y análisis de causas",
        "Excel avanzado y tablas dinámicas",
        "Power BI, bases de datos y Python",
        "Termodinámica y balances de materia",
    ]:
        add_sidebar_bullet(doc, item)

    add_heading(doc, "Educación", size=14.0, top=9, bottom=5)
    add_education(doc, "Diplomado en Ciencia de Datos", "Tecnológico Nacional de México", "2023 - 2024")
    add_education(doc, "M.Sc. en Ingeniería", "Tecnológico Nacional de México", "2021 - 2023")
    add_education(doc, "Ingeniería Química", "Universidad Pontificia Bolivariana", "2016 - 2021")

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.COLUMN)

    add_profile(doc)
    add_heading(doc, "Experiencia", size=14.2, top=5, bottom=4)

    add_job(
        doc,
        "Ingeniero de procesos y transformación digital",
        "TECONDOR S.A.S. | Contratista | Medellín",
        "2026 - Actualidad",
        [
            "Analiza y estandariza procesos de producción, compras, logística, costos y soporte administrativo, definiendo controles y trazabilidad de la información.",
            "Participa en el diseño e implementación de soluciones digitales con validaciones, cálculos automáticos, indicadores y reportes para apoyar la operación.",
            "Coordina requerimientos entre áreas y ejecuta pruebas, análisis de causas y mejoras para reducir errores y reprocesos.",
        ],
    )
    add_job(
        doc,
        "Jefe de planta",
        "PROLAQUIM S.A.S.",
        "Ene 2025 - Mar 2026",
        [
            "Lideró procesos productivos y personal operativo, asegurando el cumplimiento de especificaciones de calidad y seguridad.",
            "Estandarizó formulaciones y controló pH, viscosidad, sólidos y condiciones de operación para mejorar la repetibilidad entre lotes.",
            "Gestionó desviaciones y acciones correctivas para reducir variabilidad, reprocesos y no conformidades.",
        ],
    )
    add_job(
        doc,
        "Analista de datos",
        "Servicios de Ingeniería en Medicina S.A. de C.V.",
        "Sep 2023 - Jun 2024",
        [
            "Analizó datos de compras y contratos para identificar patrones, riesgos y oportunidades de mejora.",
            "Diseñó reportes y tableros para el seguimiento de indicadores y la toma de decisiones.",
        ],
    )
    add_job(
        doc,
        "Practicante de ingeniería de procesos",
        "Cementos Argos S.A.S.",
        "Sep 2020 - Mar 2021",
        [
            "Analizó datos operativos y automatizó reportes de producción en Excel.",
            "Apoyó el control de calidad, el análisis de no conformidades y la optimización de procesos.",
        ],
    )

    doc.core_properties.title = "Hoja de vida - Mateo Agudelo"
    doc.core_properties.subject = "Ingeniero químico - procesos, producción, calidad y mejora continua"
    doc.core_properties.author = "Mateo Agudelo"
    doc.core_properties.keywords = "ingeniero químico, procesos, producción, calidad, mejora continua, Excel, Power BI"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
